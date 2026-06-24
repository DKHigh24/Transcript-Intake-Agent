"""
build_speaker_notes.py
Generates a Word document of speaker notes / talking points for the
AI Transcript Intake Agent demo deck (19 slides).

Output: demo/AI_Transcript_Intake_Agent_Speaker_Notes.docx

Run:
    .venv/Scripts/python.exe demo/build_speaker_notes.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).resolve().parent / "AI_Transcript_Intake_Agent_Speaker_Notes.docx"

NAVY = RGBColor(0x1F, 0x38, 0x64)
TEAL = RGBColor(0x2C, 0x7D, 0x8C)
GOLD = RGBColor(0xB0, 0x7A, 0x10)
GREY = RGBColor(0x55, 0x5F, 0x6B)


# Each slide: (number, title, time, key_message, talking_points[], transition)
SLIDES = [
    (
        1, "Title - AI Transcript Intake Agent", "0:30",
        "Set the frame: this is about turning meeting conversation into tracked execution.",
        [
            "Welcome everyone. Today I'm walking you through the AI Transcript Intake Agent - a tool I built for our Electronics AI Working Group.",
            "The one-line version: it takes the transcript of our weekly meeting and turns it into structured, reviewable AI-opportunity records - and then tracks those opportunities all the way through to who is actually working on them.",
            "I'll spend about 20-25 minutes: why I built it, how it works, a live demo, the time savings, and how any of your teams could use the same engine.",
        ],
        "Let's start with the agenda so you know where we're headed.",
    ),
    (
        2, "Agenda", "0:30",
        "Give the audience the roadmap so they can relax and follow along.",
        [
            "Four parts. First, the problem - why this exists. Second, how it actually works under the hood. Third, a live demo. Fourth, the benefits and how it scales to other teams.",
            "Hold your questions if you can until the demo and the scale section - most things will get answered along the way, but I'll leave plenty of time at the end.",
        ],
        "Let's start with the problem it solves.",
    ),
    (
        3, "Part 1 - Why we built it", "0:10",
        "Section transition.",
        [
            "Part one: why this was built in the first place.",
        ],
        "",
    ),
    (
        4, "The Problem", "2:00",
        "The pain was real, weekly, and recurring - and ideas died after the meeting.",
        [
            "Every week our group meets and talks about ways the business could use AI. Great conversations - but the ideas were getting lost.",
            "Three problems. ONE - it was manual and slow. Someone had to read the entire transcript, pull out the real opportunities by hand, and fill in about 40 fields per idea in a spreadsheet. That was roughly seven hours of work, every single week.",
            "TWO - it was inconsistent. Different people captured different things, used different language, and there was no shared way to classify or compare ideas week over week.",
            "THREE - and this is the big one - there was no follow-through. We'd capture an idea, and then it would just... sit there. Nobody knew if anyone had picked it up. 'Whatever happened to that invoice-automation idea from three weeks ago?' - blank stares.",
            "The core insight that drove the whole design is at the bottom: a transcript is the BEGINNING of the work, not the end of it. Capturing the idea is step one, not the finish line.",
        ],
        "So here's what I built to fix that.",
    ),
    (
        5, "The Solution", "2:00",
        "A local agent: transcript in, reviewable opportunities + reports + tracked work out - human-governed throughout.",
        [
            "The solution is a local agent. You drop in a meeting transcript, and you get back reviewable AI-opportunity records, trend reports, and tracked work items - with a human in the loop at every step.",
            "Across the top is the flow: Read the transcript, Extract the opportunities, Classify them against our framework, Report on them weekly and monthly, and Track them in Azure DevOps.",
            "Two things make it trustworthy - left card. Full transcripts are NEVER sent to the AI model; only keyword-filtered chunks. And every AI output is draft-only - it defaults to 'Needs Review.' A human approves before anything goes anywhere official.",
            "Two things make it efficient - right card. Deterministic Python does all the heavy lifting, which is free and reproducible. The AI is used only for the two steps where judgement is actually needed. And it runs on our existing GitHub Copilot subscription - no new API key, no new spend.",
        ],
        "Let me open the hood and show you how it actually works.",
    ),
    (
        6, "Part 2 - How it works", "0:10",
        "Section transition.",
        [
            "Part two: how it works. I'll keep this practical, not deeply technical.",
        ],
        "",
    ),
    (
        7, "How It Works - The Pipeline", "2:30",
        "It's a staged pipeline; most steps are plain code, AI only where it adds value.",
        [
            "This is the whole pipeline, top to bottom. The key thing to notice is the 'Type' column on the right - most steps are plain Python, and only two are AI.",
            "Step 0 - before anything else, it syncs live status from Azure DevOps so the reports always reflect reality. I'll come back to that.",
            "Steps 1 to 3 - read the Word doc, clean it, split it into speaker turns, and keyword-filter down to just the relevant chunks. All deterministic.",
            "Steps 4 and 5 - the gold rows - are the AI. Four extracts the actual opportunities from the filtered chunks. Five classifies each one against our 40-field framework.",
            "Step 5b validates everything against allowed values. Steps 6 through 11 export the workbook, the payload, the HTML reports, archive the week, and update history.",
            "And Step 9c - opt-in - pushes approved opportunities into Azure DevOps as tracked work items. That only runs when I explicitly ask for it.",
            "The takeaway: the AI footprint is small and controlled. Everything around it is testable, reproducible code.",
        ],
        "Let me make that AI-versus-code split crystal clear.",
    ),
    (
        8, "How It Works - AI vs. Deterministic", "1:30",
        "Two AI steps for judgement; everything else is deterministic for reproducibility.",
        [
            "Left side - what the AI does. Just two things. It reads the filtered chunks and asks 'is this a real AI opportunity?' Then for each one it fills in the framework fields. Small, cheap prompts - one call per chunk or per candidate. Why AI here? Because this needs judgement and language understanding that rules simply can't capture.",
            "Right side - what Python does. Everything else: reading the document, cleaning, chunking, validation, Excel, HTML reports, history, trend analysis, deduplication, and the Azure DevOps API calls. Why Python here? Because it's reproducible, free, fast, and testable.",
            "This split is deliberate. It keeps cost down, keeps behavior predictable, and means the AI is only ever doing the part humans are worst at - reading hours of text consistently.",
        ],
        "Now, what is that framework the AI classifies against?",
    ),
    (
        9, "The Classification Framework", "2:00",
        "A shared language makes opportunities comparable, prioritizable, and trackable - and it's just config.",
        [
            "Every opportunity gets mapped onto these consistent dimensions. This is what gives us a shared language.",
            "AI Use Case Type - is this AI to understand the work, or to act on it? Operating Bucket - which part of the business: pre-sale, manufacturing, post-shipment, governance. Level of Analysis - a maturity ladder from a raw signal up to a release candidate.",
            "Maturity Signal - is this aspirational, being explored, piloting, or already deployed? Signal Strength - is it one isolated example or a cross-functional, leadership-level priority? And Value, Effort, and Risk scores for prioritization.",
            "Here's the part that matters for scaling - and the line at the bottom: this framework is just configuration. It's a JSON file. Swap it out, and the exact same engine will classify a completely different domain. Hold that thought - I'll come back to it.",
        ],
        "So what do you actually get out of it?",
    ),
    (
        10, "The Outputs", "2:00",
        "Self-contained, shareable outputs for review, leadership, Power BI, and the meeting itself.",
        [
            "Everything it produces is self-contained - HTML you open in a browser, no server needed.",
            "The weekly report has five tabs: Cards - one per opportunity with a live status chip; Analytics - charts; the Full Table - all 40 fields; Trends - what's new versus carried over; and Progress - live status from Azure DevOps. I'll show these live in a minute.",
            "The monthly rollup gives leadership the bigger picture - unique counts, what's new this month, momentum, escalations.",
            "For review and hand-off, there's an Excel workbook, a draft SharePoint payload, and a master file that feeds Power BI - including the Azure DevOps columns.",
            "And for the meeting itself, it even auto-generates the PowerPoint deck for the next session, built from the accumulated history rather than by hand.",
        ],
        "Now the newest and, honestly, the most important capability - closing the loop.",
    ),
    (
        11, "Closing the Loop - Azure DevOps", "2:30",
        "This is what turns a meeting note into a managed pipeline of work.",
        [
            "This is the capability I'm most excited about, and it's the answer to that 'whatever happened to that idea?' problem.",
            "Three pieces. PUSH - opt-in. After review, I run it with a --push-ado flag, and each approved opportunity becomes a tracked Issue in Azure DevOps, with a full description - the problem, the evidence quote, the next step. It even de-duplicates by title, so it's safe to re-run.",
            "SYNC - automatic. Every weekly run starts by pulling live status back from Azure DevOps - To Do, Doing, Done, who it's assigned to. Important: this is read-only. It never creates or edits work items. It just reads.",
            "SEE IT - the Progress tab. Every opportunity grouped by its current state, with items that moved in the last seven days highlighted, and direct links into Azure DevOps. That tab answers 'where are we at with X?' in one glance.",
            "The line at the bottom is the whole point: this is the difference between a meeting note and a managed pipeline of work. We added this specifically to close the loop on where the work actually happens.",
        ],
        "Enough slides - let me show you the real thing.",
    ),
    (
        12, "Part 3 - See it in action", "0:10",
        "Section transition into the live demo.",
        [
            "Part three - let's see it run.",
        ],
        "",
    ),
    (
        13, "Live Demo", "5:00",
        "Show, don't tell: run the pipeline, walk the report, push to ADO, show Progress.",
        [
            "DEMO SCRIPT - follow these six steps live:",
            "1. Show the input folder. Drop a DOCX transcript into input/transcripts/. Point out that the date is in the filename.",
            "2. Run the command: python src/main.py --input \"...\" --mode weekly. Narrate the steps as they scroll - read, chunk, extract, classify.",
            "3. Open the weekly report. Walk all five tabs: Cards, Analytics, Full Table, Trends, and Progress. Spend the most time on Cards and Progress.",
            "4. Click into a couple of cards. Point out that everything says 'Needs Review' - the human is in control.",
            "5. Re-run with --push-ado. Switch to the Azure DevOps board and show the new Issues appearing under the parent Epic.",
            "6. Refresh the Progress tab - show the items now linked to live ADO status.",
            "FALLBACK: if the live run is slow or the network misbehaves, I have a pre-generated report open in another tab to walk through instead.",
            "Keep narrating what you're doing and WHY - the audience cares about the workflow, not the console output.",
        ],
        "Let me bring it back to why this matters for the business.",
    ),
    (
        14, "Part 4 - Benefits & scale", "0:10",
        "Section transition.",
        [
            "Part four - the benefits, and how this scales beyond our group.",
        ],
        "",
    ),
    (
        15, "The Benefits - Time & ROI", "2:30",
        "Roughly 7.5 hours/week saved; the numbers are conservative and measurable.",
        [
            "Let's talk numbers. The headline: about 7.5 hours saved every week, an 83% reduction in manual effort, roughly 390 hours a year. At a loaded rate that's about $33,000 a year in time value - and that's a lens, not a billing claim.",
            "It breaks into two buckets. The transcript processing itself saves about 5.8 hours a week - reading, identifying, classifying, and building the reports all collapse from hours to minutes.",
            "The Azure DevOps lifecycle adds another ~1.7 hours a week - work-item creation drops from a couple of hours per cycle to minutes, status tracking goes from an hour a week to about five minutes, and traceability lookups basically disappear.",
            "Two things I want to be honest about: QA and human review time is largely RETAINED - that's the judgement step and we want to keep it. And the ADO piece consumes no AI credits at all - it's a plain REST API.",
            "These estimates are deliberately conservative. If anything, the consistency benefit - everyone speaking the same language - is worth more than the raw hours.",
        ],
        "And here's why this isn't just for our team.",
    ),
    (
        16, "How It Scales", "2:30",
        "The engine is domain-agnostic - swap config + prompt, point at new sources, serve any team.",
        [
            "This is the part I want every team lead in the room to hear. The engine is domain-agnostic. Remember the framework is just config, and the extraction logic is just a short prompt file. Swap those two things, point it at different documents, and it works for a completely different team.",
            "Sales or CX - run it on call transcripts to extract customer pain points, feature requests, and churn signals.",
            "HR or People - interview and survey notes become themes, sentiment, and training needs.",
            "Engineering - incident retros become action items, root causes, and recurring failure modes.",
            "Product - user research sessions become jobs-to-be-done and prioritized asks. Legal - contract reviews surface obligations and non-standard clauses. Operations - standups surface blockers and dependencies.",
            "And the line at the bottom is key: everything that makes this safe and useful stays the same - the ingestion, the draft-only guardrails, the human review, the trend reporting, and the Azure DevOps loop. You're only swapping the domain knowledge, not rebuilding the machine.",
        ],
        "So if you're thinking 'I want this,' here's what it takes.",
    ),
    (
        17, "Adopting It For Your Team", "1:30",
        "Five concrete steps; low barrier; no new infrastructure.",
        [
            "Five steps to stand up a new instance. One - define your framework: the dimensions and choice values for your domain, in a config file. Two - write the extraction prompt: one short file describing what an 'opportunity' looks like for you.",
            "Three - point it at your sources: drop your transcripts, notes, or documents into the input folder. Four - run and review: the same pipeline produces drafts, and a human reviews them. Five - wire your tracker: map the outputs to Azure DevOps, Jira, SharePoint - whatever closes your loop.",
            "The bottom line in green: no new infrastructure. It runs locally on an existing Copilot subscription, and it's governed by human review end to end. The barrier to try it is genuinely low.",
        ],
        "Let me wrap up.",
    ),
    (
        18, "Recap & Next Steps", "1:30",
        "Restate the value and make the ask.",
        [
            "Quick recap. It turns weekly meeting talk into structured, classified, reviewable opportunities. It saves roughly 7.5 hours a week. It closes the loop with Azure DevOps, so discussion becomes tracked execution. It's human-in-the-loop and draft-only at every step - safe by design. And it's domain-agnostic - swap the config and it serves any team.",
            "The ask: if any of this resonated for your team's workflow, come find me. I'd love to help you stand up an instance - the engine already exists; we just point it at your world.",
            "That's David High - happy to follow up with anyone after this.",
        ],
        "And with that - thank you.",
    ),
    (
        19, "Thank You / Q&A", "Remaining",
        "Open the floor.",
        [
            "Thank you. I'll open it up for questions and discussion.",
            "ANTICIPATED Q&A:",
            "- 'Is our data safe?' Full transcripts never go to the model; only filtered chunks. Everything is draft and human-reviewed. Runs locally.",
            "- 'What does it cost?' Roughly 30 Copilot credits per transcript on our existing subscription. No new license. ADO uses no credits.",
            "- 'What if the AI gets it wrong?' Everything defaults to 'Needs Review.' Nothing is published or pushed without a human approving it.",
            "- 'How long to set up for my team?' If you can describe your framework and what an opportunity looks like, we can have a first run in days, not weeks.",
            "- 'Does it have to be Azure DevOps?' No. ADO is what we use; the same outputs can map to Jira, SharePoint, or anything else.",
        ],
        "",
    ),
]


def add_heading_block(doc, slide_num, title, time):
    p = doc.add_paragraph()
    run = p.add_run(f"Slide {slide_num}  -  {title}")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = NAVY
    p.space_before = Pt(10)
    p.space_after = Pt(2)

    pt = doc.add_paragraph()
    r = pt.add_run(f"Suggested time: {time}")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    pt.space_after = Pt(4)


def add_key_message(doc, msg):
    p = doc.add_paragraph()
    label = p.add_run("Key message:  ")
    label.bold = True
    label.font.size = Pt(10.5)
    label.font.color.rgb = TEAL
    body = p.add_run(msg)
    body.font.size = Pt(10.5)
    body.font.color.rgb = GREY
    p.space_after = Pt(6)


def add_talking_points(doc, points):
    p = doc.add_paragraph()
    r = p.add_run("Talking points")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = NAVY
    p.space_after = Pt(2)
    for pt_text in points:
        bp = doc.add_paragraph(style="List Bullet")
        run = bp.add_run(pt_text)
        run.font.size = Pt(11)
        bp.space_after = Pt(3)


def add_transition(doc, text):
    if not text:
        return
    p = doc.add_paragraph()
    label = p.add_run("Transition:  ")
    label.bold = True
    label.italic = True
    label.font.size = Pt(10)
    label.font.color.rgb = GOLD
    body = p.add_run(text)
    body.italic = True
    body.font.size = Pt(10)
    body.font.color.rgb = GOLD
    p.space_before = Pt(2)
    p.space_after = Pt(8)


def add_divider(doc):
    p = doc.add_paragraph()
    r = p.add_run("_" * 60)
    r.font.color.rgb = RGBColor(0xDD, 0xE3, 0xED)
    r.font.size = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(2)


def build():
    doc = Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    # Title
    t = doc.add_paragraph()
    tr = t.add_run("AI Transcript Intake Agent")
    tr.bold = True
    tr.font.size = Pt(24)
    tr.font.color.rgb = NAVY
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.space_after = Pt(0)

    s = doc.add_paragraph()
    sr = s.add_run("Speaker Notes & Talking Points  -  Demo + Overview (20-30 min)")
    sr.font.size = Pt(13)
    sr.font.color.rgb = TEAL
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s.space_after = Pt(4)

    meta = doc.add_paragraph()
    mr = meta.add_run("Electronics AI Working Group  |  Presenter: David High (DK High)")
    mr.font.size = Pt(10)
    mr.font.color.rgb = GREY
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.space_after = Pt(6)

    # How-to-use note
    how = doc.add_paragraph()
    hr = how.add_run(
        "How to use this document:  Each slide below has a key message, bulleted "
        "talking points, and a one-line transition to the next slide. Copy the "
        "talking points into the Notes pane of the matching slide in "
        "AI_Transcript_Intake_Agent_Demo.pptx. Times are a guide totaling ~25 minutes "
        "plus Q&A - adjust to your slot."
    )
    hr.italic = True
    hr.font.size = Pt(10)
    hr.font.color.rgb = GREY
    how.space_after = Pt(8)

    # Total time line
    tot = doc.add_paragraph()
    totr = tot.add_run(
        "Pacing summary:  Parts 1-2 ~ 13 min  |  Demo ~ 5 min  |  Part 4 ~ 8 min  |  Q&A = remaining"
    )
    totr.bold = True
    totr.font.size = Pt(10)
    totr.font.color.rgb = NAVY
    tot.space_after = Pt(10)

    add_divider(doc)

    for (num, title, time, key, points, transition) in SLIDES:
        add_heading_block(doc, num, title, time)
        add_key_message(doc, key)
        add_talking_points(doc, points)
        add_transition(doc, transition)
        add_divider(doc)

    # Closing tips page
    doc.add_page_break()
    ct = doc.add_paragraph()
    ctr = ct.add_run("Delivery Tips")
    ctr.bold = True
    ctr.font.size = Pt(16)
    ctr.font.color.rgb = NAVY
    ct.space_after = Pt(6)

    tips = [
        "Open strong with the pain everyone recognizes: 'whatever happened to that idea?' Get heads nodding before any tech.",
        "Repeat the core line twice: 'a transcript is the beginning of the work, not the end.' It frames the whole talk.",
        "On the pipeline slide, resist going too deep - point at the 'Type' column and move on. Depth lives in the demo.",
        "Make the demo the centerpiece. If you only have 15 minutes, cut slides, not the demo.",
        "Always have a pre-generated report open as a fallback in case the live run is slow.",
        "On the scale slide, look at specific people: 'Sarah, for your CX calls...' Make it personal.",
        "Close with a clear, low-friction ask: 'come find me and we'll point it at your world.'",
        "Keep the guardrail message visible throughout: draft-only, human-in-the-loop, runs locally. It preempts the data-safety question.",
    ]
    for tip in tips:
        bp = doc.add_paragraph(style="List Bullet")
        r = bp.add_run(tip)
        r.font.size = Pt(11)
        bp.space_after = Pt(4)

    doc.save(str(OUT))
    print(f"Saved {OUT}")
    print(f"Paragraphs: {len(doc.paragraphs)}")


if __name__ == "__main__":
    build()
